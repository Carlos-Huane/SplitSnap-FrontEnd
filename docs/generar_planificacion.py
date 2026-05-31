"""
Genera el documento Word de planificación del Sprint Frontend de SplitSnap.
Salida: Planificacion-Sprint-Frontend.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def h1(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)


def h2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)


def h3(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)


def bullets(doc, items):
    for it in items:
        doc.add_paragraph(it, style='List Bullet')


def numbered(doc, items):
    for it in items:
        doc.add_paragraph(it, style='List Number')


def hu_block(doc, code, title, role, want, so_that, criteria, tasks, sp, owner, depends=None):
    h3(doc, f"{code} — {title}")

    p = doc.add_paragraph()
    p.add_run("Asignada a: ").bold = True
    p.add_run(owner)

    p = doc.add_paragraph()
    p.add_run("Story Points: ").bold = True
    p.add_run(str(sp))

    if depends:
        p = doc.add_paragraph()
        p.add_run("Depende de: ").bold = True
        p.add_run(depends)

    p = doc.add_paragraph()
    p.add_run("Historia de Usuario").bold = True
    p = doc.add_paragraph()
    p.add_run(f"Como {role}, quiero {want}, para {so_that}.")

    p = doc.add_paragraph()
    p.add_run("Criterios de aceptación").bold = True
    bullets(doc, criteria)

    p = doc.add_paragraph()
    p.add_run("Tareas técnicas (subtasks Jira)").bold = True
    numbered(doc, tasks)

    doc.add_paragraph()


def main():
    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)

    # PORTADA
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SplitSnap — Sprint Integración Frontend")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("Planificación de Historias de Usuario y Asignaciones")
    run.font.size = Pt(14)
    run.italic = True

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run("Curso: Herramientas de Desarrollo  |  Equipo: 5 integrantes  |  Duración: 2 semanas")
    run.font.size = Pt(11)

    doc.add_paragraph()

    # OBJETIVO
    h1(doc, "Objetivo del Sprint")
    doc.add_paragraph(
        "Reemplazar los datos mock del AppContext + localStorage por llamadas reales al backend desplegado, "
        "dejando la aplicación lista para la exposición y evaluación del docente. Cada integrante toma una épica "
        "que corresponde a un conjunto coherente de pantallas y endpoints."
    )

    # CONTEXTO TECNICO
    h1(doc, "Contexto técnico")
    bullets(doc, [
        "Backend desplegado en Railway: la URL pública será compartida por el líder técnico (Carlos).",
        "Cada integrante setea VITE_API_URL=<url-pública> en su archivo .env local del frontend.",
        "Autenticación por JWT: el token se guarda en localStorage y se envía en el header Authorization.",
        "Las llamadas HTTP se hacen a través del cliente axios centralizado (entregable de la épica F0).",
        "Frontend: React 19 + Vite + react-router-dom. Estado global en AppContext.",
        "El OCR usa Google Cloud Vision (free tier 1000 imágenes/mes, cubierto). El JSON del service account está solo en el servidor.",
    ])

    # RESUMEN DE ASIGNACIONES
    h1(doc, "Resumen de asignaciones")
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    headers = ["Épica", "Responsable", "Pantallas principales", "Story Points", "Prioridad"]
    for i, h in enumerate(headers):
        hdr[i].text = h
        for r in hdr[i].paragraphs[0].runs:
            r.bold = True
        set_cell_bg(hdr[i], "1F4E79")
        for r in hdr[i].paragraphs[0].runs:
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    rows_data = [
        ("F0 — Infraestructura compartida", "Carlos Huane",
         "Cliente axios + Auth global", "5", "🔴 Crítica"),
        ("F1 — Autenticación y Perfil", "Marcela",
         "Register, Login, Profile, Créditos", "13", "🔴 Crítica"),
        ("F2 — Grupos: listado/crear/invitar", "Obed Velarde",
         "Dashboard, GroupList, CreateGroup, InviteMembers", "13", "🟡 Alta"),
        ("F3 — Detalle de Grupo y Gastos Manuales", "Yorma Campos",
         "GroupDetail, AddExpense", "13", "🟡 Alta"),
        ("F4 — OCR de Recibos", "Carlos Huane",
         "ScanReceipt, ReviewItems", "8", "🟢 Media"),
        ("F5 — Deudas y Pagos", "Dafne Fuentes",
         "DebtSummary (pendientes + pagadas, mark-paid, pay-credits)", "13", "🟢 Media"),
    ]
    for row in rows_data:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Total: 16 historias de usuario · 65 story points").bold = True

    # CRONOGRAMA
    h1(doc, "Cronograma sugerido (2 semanas)")
    bullets(doc, [
        "Días 1-2: Carlos termina F0 (infraestructura compartida). Bloquea al resto del equipo.",
        "Días 2-4: Marcela arranca F1 (Auth) tras avisar de F0. Resto del equipo prepara su entorno local.",
        "Días 4-7: F2 (Obed) y F3 (Yorma) en paralelo. Carlos arranca F4.",
        "Días 7-10: Dafne arranca F5. Carlos termina F4 y soporta a integraciones.",
        "Días 10-12: Integración cruzada, testing, fix de bugs reportados.",
        "Días 13-14: Polish visual, deploy del frontend a Vercel, ensayo de exposición.",
    ])

    doc.add_page_break()

    # ============================================================
    # ÉPICA F0
    # ============================================================
    h1(doc, "ÉPICA F0 — Infraestructura compartida")

    p = doc.add_paragraph()
    p.add_run("Responsable: ").bold = True
    p.add_run("Carlos Huane")
    p = doc.add_paragraph()
    p.add_run("Prioridad: ").bold = True
    p.add_run("🔴 Crítica (bloquea al resto del equipo)")

    doc.add_paragraph(
        "Objetivo: dejar listo el cliente HTTP, la gestión de JWT, la estructura de servicios "
        "y el manejo global de errores 401/403/500 antes de que nadie integre sus pantallas."
    )

    hu_block(
        doc,
        code="HU-F0.1",
        title="Cliente HTTP con JWT y manejo de auth global",
        role="equipo de desarrollo",
        want="un cliente axios centralizado con interceptores que manejen el JWT y los errores 401",
        so_that="ningún integrante repita lógica de autenticación en cada request",
        criteria=[
            "Axios instalado en package.json.",
            "src/services/api.js exporta una instancia axios con baseURL = import.meta.env.VITE_API_URL (fallback http://localhost:8080).",
            "Interceptor request: lee localStorage.getItem('splitsnap_token') y agrega Authorization: Bearer <token> si existe.",
            "Interceptor response: si 401 → limpia token + dispatch LOGOUT + redirige a /login.",
            "Si 403 → muestra mensaje 'No tienes acceso a este recurso'.",
            "Si red caída → muestra mensaje 'Error de conexión, verifica el backend'.",
            ".env.example actualizado con VITE_API_URL=http://localhost:8080.",
            "RequireAuth.jsx valida la existencia del token, no solo currentUserId del context.",
        ],
        tasks=[
            "npm install axios",
            "Crear src/services/api.js con instancia + interceptores.",
            "Crear src/services/auth.service.js (login, register, logout, getToken, setToken, getCurrentUser).",
            "Refactorizar AppContext.jsx para que el token viva en localStorage.",
            "Actualizar RequireAuth.jsx.",
            "Crear .env.example y documentar en README cómo levantar front + back juntos.",
            "Crear archivos placeholder por dominio en src/services/: groups.service.js, expenses.service.js, debts.service.js, credits.service.js, users.service.js.",
        ],
        sp=5,
        owner="Carlos Huane",
    )

    doc.add_page_break()

    # ============================================================
    # ÉPICA F1
    # ============================================================
    h1(doc, "ÉPICA F1 — Autenticación y Perfil")

    p = doc.add_paragraph()
    p.add_run("Responsable: ").bold = True
    p.add_run("Marcela")
    p = doc.add_paragraph()
    p.add_run("Prioridad: ").bold = True
    p.add_run("🔴 Crítica (segunda después de F0)")
    p = doc.add_paragraph()
    p.add_run("Endpoints involucrados: ").bold = True
    p.add_run("POST /api/auth/register, POST /api/auth/login, GET/PUT /api/users/me, PUT /api/users/me/avatar, GET /api/users/me/credits, POST /api/users/me/credits/buy")

    hu_block(
        doc,
        code="HU-F1.1",
        title="Registrar nuevo usuario",
        role="visitante",
        want="crear una cuenta desde la pantalla /register con mis datos",
        so_that="acceder a SplitSnap y comenzar a usar los grupos",
        criteria=[
            "Register.jsx envía POST /api/auth/register con { name, email, phone, password }.",
            "Si éxito → guarda token en localStorage, dispatch LOGIN al context, navega a /dashboard.",
            "Si email duplicado (409) → muestra mensaje 'Este email ya está registrado'.",
            "Si validación falla (400) → muestra mensajes del backend.",
            "Loading state visible mientras espera respuesta.",
        ],
        tasks=[
            "Reemplazar mock REGISTER_USER por llamada a authService.register.",
            "Implementar useState para loading y error.",
            "Validación cliente: email formato válido, password mínimo 6 caracteres.",
            "Manejo de respuesta exitosa (guardar token + navegar).",
            "Manejo de errores (mostrar mensajes según código HTTP).",
        ],
        sp=3,
        owner="Marcela",
        depends="HU-F0.1",
    )

    hu_block(
        doc,
        code="HU-F1.2",
        title="Iniciar sesión",
        role="usuario registrado",
        want="iniciar sesión desde /login con mi email y contraseña",
        so_that="acceder a mis grupos y datos",
        criteria=[
            "Login.jsx envía POST /api/auth/login con { email, password }.",
            "Guarda token, dispatch LOGIN, navega a /dashboard.",
            "Si 401 → muestra 'Credenciales incorrectas'.",
            "Soporte de Enter para enviar el formulario.",
            "Link a /register visible para crear cuenta.",
        ],
        tasks=[
            "Implementar authService.login en services.",
            "Refactorizar Login.jsx para llamar al servicio.",
            "Estados loading y error.",
            "Persistir sesión: si hay token válido en localStorage al cargar /login, redirigir a /dashboard.",
        ],
        sp=3,
        owner="Marcela",
        depends="HU-F0.1",
    )

    hu_block(
        doc,
        code="HU-F1.3",
        title="Ver y editar perfil",
        role="usuario",
        want="ver y editar mi perfil desde /profile",
        so_that="mantener actualizados mi nombre, email, teléfono y contraseña",
        criteria=[
            "Al montar Profile.jsx → llama GET /api/users/me y carga datos en el formulario.",
            "Botón 'Editar' abre el formulario en modo edición.",
            "Cambios se envían vía PUT /api/users/me.",
            "Cambio de password requiere campo currentPassword.",
            "Subir avatar → PUT /api/users/me/avatar con multipart/form-data, muestra preview.",
            "Validación de archivo: solo imágenes (jpg, png, webp), tamaño máximo 2MB.",
        ],
        tasks=[
            "Implementar usersService.getMe(), updateMe(payload), uploadAvatar(file).",
            "Componente ProfileForm con campos editables.",
            "Componente AvatarUploader con preview y validación.",
            "Manejo de errores 400 (validación) y 409 (email duplicado).",
        ],
        sp=4,
        owner="Marcela",
        depends="HU-F0.1, HU-F1.2",
    )

    hu_block(
        doc,
        code="HU-F1.4",
        title="Ver y comprar créditos",
        role="usuario",
        want="ver mi balance de créditos y comprar más desde mi perfil",
        so_that="poder pagar deudas con créditos del sistema sin salir de la app",
        criteria=[
            "Sección de créditos en Profile.jsx muestra balance + historial de transacciones (PURCHASE y SPEND).",
            "Al montar → GET /api/users/me/credits.",
            "Botón 'Comprar créditos' abre modal con paquetes 10, 25, 50, 100 (o monto personalizado > 0).",
            "Al confirmar compra → POST /api/users/me/credits/buy con { amount }.",
            "Tras compra → refresca balance y muestra mensaje de éxito.",
        ],
        tasks=[
            "Implementar creditsService.getCredits(), buyCredits(amount).",
            "Componente CreditsPanel con balance y tabla de transacciones.",
            "Modal BuyCreditsModal con selector de paquetes.",
            "Reemplazar BUY_CREDITS y SPEND_CREDITS del context por datos del backend.",
        ],
        sp=3,
        owner="Marcela",
        depends="HU-F0.1, HU-F1.2",
    )

    doc.add_page_break()

    # ============================================================
    # ÉPICA F2
    # ============================================================
    h1(doc, "ÉPICA F2 — Grupos: Listado, Creación e Invitaciones")

    p = doc.add_paragraph()
    p.add_run("Responsable: ").bold = True
    p.add_run("Obed Velarde")
    p = doc.add_paragraph()
    p.add_run("Prioridad: ").bold = True
    p.add_run("🟡 Alta")
    p = doc.add_paragraph()
    p.add_run("Endpoints involucrados: ").bold = True
    p.add_run("GET /api/groups, POST /api/groups, POST /api/groups/{id}/members, GET /api/users/search")

    doc.add_paragraph(
        "⚠️ Gap conocido del backend: GET /api/groups no devuelve el campo myBalance todavía. "
        "Mostrar 'S/ 0.00' como placeholder o calcular del lado del front sumando deudas pendientes."
    )

    hu_block(
        doc,
        code="HU-F2.1",
        title="Listar mis grupos en Dashboard",
        role="usuario",
        want="ver mis grupos en /dashboard y /groups",
        so_that="poder elegir uno y revisar sus gastos y deudas",
        criteria=[
            "Al montar Dashboard.jsx y GroupList.jsx → GET /api/groups.",
            "Cada tarjeta muestra nombre, emoji, cantidad de miembros.",
            "Si la lista está vacía → navegar automáticamente a /empty.",
            "Loading state mientras espera respuesta.",
            "Si error de red → mostrar mensaje + botón 'Reintentar'.",
        ],
        tasks=[
            "Implementar groupsService.getMyGroups().",
            "Refactorizar Dashboard.jsx y GroupList.jsx para usar el servicio.",
            "Eliminar la dependencia del array groups del AppContext.",
            "Calcular myBalance del lado del front si el backend no lo devuelve.",
        ],
        sp=4,
        owner="Obed Velarde",
        depends="HU-F0.1, HU-F1.2",
    )

    hu_block(
        doc,
        code="HU-F2.2",
        title="Crear grupo nuevo",
        role="usuario",
        want="crear un grupo desde /groups/new con nombre, emoji y miembros iniciales",
        so_that="comenzar a registrar gastos compartidos",
        criteria=[
            "CreateGroup.jsx envía POST /api/groups con { name, emoji, memberIds }.",
            "Validación cliente: nombre no vacío.",
            "Emoji opcional (default: 📦).",
            "Tras crear → navegar a /groups/{id} (detalle del nuevo grupo).",
            "Si error → mostrar mensaje y mantener formulario.",
        ],
        tasks=[
            "Implementar groupsService.createGroup(payload).",
            "Reemplazar mock ADD_GROUP por llamada al servicio.",
            "Mantener selector de emojis del front actual.",
            "Manejo de errores y loading.",
        ],
        sp=3,
        owner="Obed Velarde",
        depends="HU-F0.1, HU-F2.1",
    )

    hu_block(
        doc,
        code="HU-F2.3",
        title="Invitar miembros al grupo",
        role="miembro del grupo",
        want="agregar más personas al grupo desde /groups/:id/invite",
        so_that="incluirlas en los gastos compartidos",
        criteria=[
            "Input de búsqueda → GET /api/users/search?q={query} con debounce de 300ms.",
            "Mínimo 2 caracteres para activar la búsqueda.",
            "Resultados muestran nombre, email, avatar.",
            "Botón 'Agregar' por usuario → POST /api/groups/{id}/members con { userId }.",
            "Si usuario ya es miembro (409) → mostrar 'Ya está en el grupo'.",
            "Tras invitar → actualizar lista y mantener el modal abierto para invitar a más.",
        ],
        tasks=[
            "Implementar usersService.search(query) y groupsService.addMember(groupId, userId).",
            "Componente UserSearchInput con debounce.",
            "Manejo de estados: empty, loading, results, error.",
            "Reemplazar mock ADD_MEMBER_TO_GROUP por llamada al servicio.",
        ],
        sp=6,
        owner="Obed Velarde",
        depends="HU-F0.1, HU-F2.2",
    )

    doc.add_page_break()

    # ============================================================
    # ÉPICA F3
    # ============================================================
    h1(doc, "ÉPICA F3 — Detalle de Grupo y Gastos Manuales")

    p = doc.add_paragraph()
    p.add_run("Responsable: ").bold = True
    p.add_run("Yorma Campos")
    p = doc.add_paragraph()
    p.add_run("Prioridad: ").bold = True
    p.add_run("🟡 Alta")
    p = doc.add_paragraph()
    p.add_run("Endpoints involucrados: ").bold = True
    p.add_run("GET /api/groups/{id}, GET /api/groups/{id}/expenses, POST /api/groups/{id}/expenses, GET /api/groups/{id}/expenses/{expenseId}")

    doc.add_paragraph(
        "⚠️ Gap conocido del backend: ExpenseResponse del listado no incluye splitBetween[]. "
        "Para mostrar el detalle de un gasto al hacer clic, llamar a GET /api/groups/{id}/expenses/{expenseId}."
    )

    hu_block(
        doc,
        code="HU-F3.1",
        title="Ver detalle de un grupo",
        role="miembro del grupo",
        want="ver /groups/:id con miembros, gastos recientes y resumen",
        so_that="entender el estado financiero del grupo",
        criteria=[
            "Al montar GroupDetail.jsx → GET /api/groups/{groupId} (datos del grupo y miembros).",
            "También llama GET /api/groups/{groupId}/expenses (lista de gastos).",
            "Muestra lista de miembros con avatares.",
            "Muestra lista de gastos recientes ordenada por fecha descendente.",
            "Si no es miembro (403) → mostrar mensaje 'No tienes acceso' + botón 'Volver'.",
            "Botones para navegar a 'Agregar gasto', 'Escanear recibo', 'Invitar miembros', 'Ver deudas'.",
        ],
        tasks=[
            "Implementar groupsService.getGroup(groupId).",
            "Implementar expensesService.getExpensesByGroup(groupId).",
            "Refactorizar GroupDetail.jsx para usar los servicios.",
            "Componente MemberList y ExpenseList.",
            "Manejo de errores 403 y 404.",
        ],
        sp=5,
        owner="Yorma Campos",
        depends="HU-F0.1, HU-F2.1",
    )

    hu_block(
        doc,
        code="HU-F3.2",
        title="Registrar gasto manual",
        role="miembro del grupo",
        want="agregar un gasto manual desde /groups/:id/add-expense",
        so_that="que el sistema calcule las deudas automáticamente",
        criteria=[
            "AddExpense.jsx envía POST /api/groups/{id}/expenses con { description, amount, paidBy, date, splitBetween[] }.",
            "Selector '¿quién pagó?' con todos los miembros (default: usuario actual).",
            "Selector 'dividir entre' con checkbox por miembro.",
            "Auto-cálculo: monto / miembros seleccionados = monto por persona (editable por persona).",
            "Validación cliente: suma de splits debe igualar amount total (tolerancia 0.01).",
            "Validación: description no vacía, amount > 0.",
            "Si éxito → navegar a /groups/:id.",
            "Si error de validación del backend → mostrar mensaje.",
        ],
        tasks=[
            "Implementar expensesService.createExpense(groupId, payload).",
            "Refactorizar AddExpense.jsx para usar el servicio.",
            "Lógica de auto-split entre miembros seleccionados.",
            "Validación cliente robusta antes de enviar.",
            "Manejo de loading y errores.",
        ],
        sp=8,
        owner="Yorma Campos",
        depends="HU-F0.1, HU-F3.1",
    )

    doc.add_page_break()

    # ============================================================
    # ÉPICA F4
    # ============================================================
    h1(doc, "ÉPICA F4 — OCR de Recibos")

    p = doc.add_paragraph()
    p.add_run("Responsable: ").bold = True
    p.add_run("Carlos Huane")
    p = doc.add_paragraph()
    p.add_run("Prioridad: ").bold = True
    p.add_run("🟢 Media")
    p = doc.add_paragraph()
    p.add_run("Endpoints involucrados: ").bold = True
    p.add_run("POST /api/ocr/scan, POST /api/groups/{id}/expenses")

    doc.add_paragraph(
        "Carlos asume esta épica porque tiene configurado el service account de Google Cloud Vision. "
        "Una vez el backend esté en Railway con el JSON cargado como variable de entorno, todo el equipo "
        "podrá probar OCR consumiendo el endpoint sin configuración adicional."
    )

    hu_block(
        doc,
        code="HU-F4.1",
        title="Escanear recibo",
        role="usuario",
        want="subir una imagen de un recibo desde /groups/:id/scan",
        so_that="que el sistema lea el monto y la descripción automáticamente",
        criteria=[
            "ScanReceipt.jsx permite seleccionar imagen con input file (accept=image/*, capture=environment para abrir cámara en mobile).",
            "Preview de la imagen antes de enviar.",
            "Botón 'Escanear' → POST /api/ocr/scan con FormData (Content-Type: multipart/form-data).",
            "Spinner con texto 'Leyendo recibo...' mientras procesa.",
            "Si éxito → navegar a /groups/:id/scan/review pasando por router state { detectedAmount, extractedItems, description }.",
            "Si 500 (Google Vision falla) → mensaje 'No pudimos leer el recibo, intenta con otra foto o ingrésalo manual'.",
            "Botón 'Ingresar manual' → navegar a /groups/:id/add-expense.",
        ],
        tasks=[
            "Implementar expensesService.scanReceipt(file) con FormData.",
            "Refactorizar ScanReceipt.jsx (eliminar mockReceiptItems).",
            "Manejar loading y errores.",
            "Pasar respuesta al siguiente paso vía router state.",
        ],
        sp=3,
        owner="Carlos Huane",
        depends="HU-F0.1, HU-F3.1",
    )

    hu_block(
        doc,
        code="HU-F4.2",
        title="Revisar e ingresar gasto desde OCR",
        role="usuario",
        want="revisar y editar lo detectado antes de confirmar el gasto",
        so_that="corregir errores del OCR antes de guardarlo",
        criteria=[
            "ReviewItems.jsx muestra description editable (prefill), amount editable (prefill con detectedAmount).",
            "Textarea o pre con el texto crudo extraído como referencia (solo lectura).",
            "Selector '¿quién pagó?' (default: usuario actual).",
            "Selector 'dividir entre' (checkbox por miembro).",
            "Auto-cálculo: monto / miembros seleccionados (editable por persona).",
            "Validación: suma de splits == amount (tolerancia 0.01).",
            "Botón 'Confirmar gasto' → POST /api/groups/{id}/expenses.",
            "Si éxito → navegar a /groups/:id.",
        ],
        tasks=[
            "Refactorizar ReviewItems.jsx para tomar datos del router state.",
            "Reutilizar la misma lógica de split que F3.2 (idealmente extraer a un hook compartido).",
            "Llamada a expensesService.createExpense reutilizando el servicio de F3.",
            "Validación y manejo de errores.",
        ],
        sp=4,
        owner="Carlos Huane",
        depends="HU-F0.1, HU-F4.1, HU-F3.2",
    )

    hu_block(
        doc,
        code="HU-F4.3",
        title="Documentar setup de Google Cloud Vision (técnica)",
        role="integrante del equipo",
        want="saber cómo el OCR funciona en el backend desplegado y cómo testearlo local si fuera necesario",
        so_that="poder mantener o debuggear el OCR sin depender de Carlos",
        criteria=[
            "Sección en el README del backend con pasos para crear cuenta GCP, habilitar Cloud Vision API, crear service account, descargar JSON, guardarlo en .gcp/, setear GOOGLE_APPLICATION_CREDENTIALS en .env.",
            "Mencionar que el free tier es 1000 imágenes/mes y que el JSON nunca debe subirse al repo.",
            "Mencionar que en producción (Railway) el JSON está como variable de entorno o secret file.",
        ],
        tasks=[
            "Redactar sección 'OCR setup' en el README.",
            "Actualizar .env.example con la variable.",
        ],
        sp=1,
        owner="Carlos Huane",
        depends="HU-F4.1",
    )

    doc.add_page_break()

    # ============================================================
    # ÉPICA F5
    # ============================================================
    h1(doc, "ÉPICA F5 — Deudas y Pagos")

    p = doc.add_paragraph()
    p.add_run("Responsable: ").bold = True
    p.add_run("Dafne Fuentes")
    p = doc.add_paragraph()
    p.add_run("Prioridad: ").bold = True
    p.add_run("🟢 Media")
    p = doc.add_paragraph()
    p.add_run("Endpoints involucrados: ").bold = True
    p.add_run("GET /api/groups/{id}/debts, PUT /api/groups/{id}/debts/{debtId}/mark-paid, PUT /api/groups/{id}/debts/{debtId}/pay-credits")

    doc.add_paragraph(
        "⚠️ Adapter del front: el backend devuelve DebtResponse con fromUser: { id, name, email, avatarUrl } (anidado). "
        "El mock del frontend usa fromUserId (flat). Mapear al recibir si es necesario, o adaptar el código del componente."
    )

    hu_block(
        doc,
        code="HU-F5.1",
        title="Ver resumen de deudas del grupo",
        role="miembro del grupo",
        want="ver /groups/:id/debts con deudas pendientes y pagadas",
        so_that="saber cuánto debo y cuánto me deben",
        criteria=[
            "Al montar DebtSummary.jsx → GET /api/groups/{groupId}/debts.",
            "Tabs 'Pendientes' y 'Pagadas' usando ?status=PENDING y ?status=PAID.",
            "Cada deuda muestra 'X debe S/ N a Y' con avatares.",
            "Cálculo de balances por miembro: suma de lo que debe + lo que le deben.",
            "Si no es miembro (403) → mensaje 'No tienes acceso'.",
        ],
        tasks=[
            "Implementar debtsService.getDebts(groupId, status).",
            "Refactorizar DebtSummary.jsx para usar el servicio.",
            "Componente DebtCard.",
            "Adapter: mapear fromUser.id → fromUserId si lo necesita el front actual.",
        ],
        sp=4,
        owner="Dafne Fuentes",
        depends="HU-F0.1, HU-F3.1",
    )

    hu_block(
        doc,
        code="HU-F5.2",
        title="Marcar deuda como pagada (manual)",
        role="usuario deudor",
        want="marcar mi deuda como pagada indicando el método de pago",
        so_that="registrar que saldé la deuda con el acreedor",
        criteria=[
            "Botón 'Marcar pagada' visible solo si fromUser.id === usuario actual.",
            "Click abre modal con selector: yape, paypal, efectivo.",
            "Confirmar → PUT /api/groups/{groupId}/debts/{debtId}/mark-paid con { paidWith }.",
            "Si éxito → refrescar lista, mover deuda a tab 'Pagadas'.",
            "Si 403 (intentó marcar deuda ajena) → mensaje 'Solo el deudor puede marcar la deuda'.",
            "Si 409 (ya estaba pagada) → mensaje 'Esta deuda ya estaba pagada' + refrescar.",
        ],
        tasks=[
            "Implementar debtsService.markAsPaid(groupId, debtId, paidWith).",
            "Componente MarkAsPaidModal con selector de método.",
            "Reemplazar mock MARK_DEBT_PAID por llamada al servicio.",
            "Manejo de errores 403 y 409.",
        ],
        sp=4,
        owner="Dafne Fuentes",
        depends="HU-F0.1, HU-F5.1",
    )

    hu_block(
        doc,
        code="HU-F5.3",
        title="Pagar deuda con créditos del sistema",
        role="deudor con créditos",
        want="saldar mi deuda usando créditos comprados",
        so_that="pagar sin salir de la aplicación",
        criteria=[
            "Botón 'Pagar con créditos' visible solo si fromUser.id === usuario actual.",
            "Mostrar balance de créditos actual antes de confirmar.",
            "Validación cliente: balance suficiente antes de llamar.",
            "Confirmar → PUT /api/groups/{groupId}/debts/{debtId}/pay-credits.",
            "Si 400 'Créditos insuficientes' → mensaje + link a /profile (sección créditos para comprar).",
            "Si éxito → actualizar balance de créditos en el context + lista de deudas.",
        ],
        tasks=[
            "Implementar debtsService.payWithCredits(groupId, debtId).",
            "Componente PayWithCreditsModal mostrando balance + monto a descontar.",
            "Refrescar balance de créditos tras pago exitoso.",
            "Manejo de errores 400, 403 y 409.",
        ],
        sp=5,
        owner="Dafne Fuentes",
        depends="HU-F0.1, HU-F5.1, HU-F1.4",
    )

    doc.add_page_break()

    # NOTAS FINALES
    h1(doc, "Notas finales y gaps conocidos")

    h2(doc, "Gaps del backend (no bloquean el sprint)")
    bullets(doc, [
        "GET /api/groups no devuelve myBalance todavía. Calcular en el front mientras tanto (épica F2).",
        "GET /api/groups/{id} no devuelve recentExpenses[] ni memberBalances[]. Hacer llamada adicional a /expenses (épica F3).",
        "GET /api/groups/{id}/expenses no incluye splitBetween[] en cada gasto del listado. Llamar al detalle individual cuando se necesite (épica F3).",
        "GET /api/users/me/transactions no soporta filtros ?from= y ?to= todavía. Si se necesitan, agregarlos al backend después.",
    ])

    h2(doc, "Pantallas fuera del scope del sprint")
    bullets(doc, [
        "Dashboard consolidado: se mantiene el dashboard básico con solo el listado de grupos (cubierto en F2).",
        "Historial global de transacciones: pantalla Historial puede quedar con datos mock por ahora.",
        "ForgotPassword: el backend no tiene endpoint, dejar la pantalla en modo 'Próximamente'.",
    ])

    h2(doc, "Definition of Done del sprint")
    bullets(doc, [
        "Todas las HUs aceptadas y mergeadas a develop.",
        "Frontend deployado en Vercel y consumiendo el backend de Railway.",
        "Login funcional con usuarios seed (carlos@splitsnap.com / test123).",
        "Flujo completo demostrable: registro → crear grupo → invitar → agregar gasto → marcar deuda pagada.",
        "OCR demostrable con al menos un recibo real.",
        "README actualizado con instrucciones de setup local y URL de producción.",
    ])

    h2(doc, "Convenciones del sprint")
    bullets(doc, [
        "Ramas: feature/F<numero>-<nombre-corto> (ej: feature/F1-auth, feature/F4-ocr).",
        "Commits: estilo Conventional Commits (feat:, fix:, refactor:, docs:, chore:).",
        "Pull Requests: hacia develop, con descripción de la HU y screenshot si aplica.",
        "Code review: mínimo 1 aprobación antes de mergear.",
    ])

    # Guardar
    out_path = "C:/Users/carlo/Desktop/hsproyect/UNIVERSIDAD/Herramientas de desarrollo/SplitSnap-FrontEnd/docs/Planificacion-Sprint-Frontend.docx"
    doc.save(out_path)
    print(f"OK: {out_path}")


if __name__ == "__main__":
    main()
